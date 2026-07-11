import os

from docx import Document


def generate_docx(
    transcript,
    summary,
    notes,
    keywords,
    quiz,
    flashcards,
    output_path="outputs/lecture_notes.docx"
):
    os.makedirs("outputs", exist_ok=True)

    document = Document()

    document.add_heading(
        "Lecture Voice-to-Notes Generator",
        level=0
    )

    document.add_paragraph(
        "AI-Generated Study Material"
    )

    sections = [
        ("Complete Transcript", transcript),
        ("Concise Summary", summary),
        ("Organized Study Notes", notes),
        ("Important Keywords", keywords),
        ("Quiz", quiz),
        ("Flashcards", flashcards)
    ]

    for heading, content in sections:

        document.add_heading(
            heading,
            level=1
        )

        document.add_paragraph(
            content or "No content generated."
        )

    document.save(output_path)

    return output_path